import streamlit as st
import fitz
from docx import Document
from fpdf import FPDF
from pdf2image import convert_from_bytes
import pytesseract
from PIL import Image
import io
import re
import time
import google.generativeai as genai

# ------------------------------------------------------------
# Configuración de la página
# ------------------------------------------------------------
st.set_page_config(page_title="Traductor de documentos", layout="wide")

# Botones de descarga más vistosos
st.markdown("""
<style>
    .stDownloadButton button {
        font-size: 18px !important;
        font-weight: bold !important;
        background-color: #FF4B4B !important;
        color: white !important;
        padding: 0.75em 2em !important;
        border-radius: 12px !important;
        border: none !important;
        box-shadow: 2px 2px 10px rgba(0,0,0,0.2);
        transition: transform 0.2s;
    }
    .stDownloadButton button:hover {
        background-color: #e04343 !important;
        transform: scale(1.05);
    }
</style>
""", unsafe_allow_html=True)

st.title("🌐 Traductor de documentos a cualquier idioma")
st.markdown("Sube un PDF, imagen o Word y obtén una traducción **natural y con formato**.")

# ------------------------------------------------------------
# Cliente de Gemini (gratuito)
# ------------------------------------------------------------
genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
modelo = genai.GenerativeModel('gemini-2.0-flash')

# ------------------------------------------------------------
# Idiomas
# ------------------------------------------------------------
IDIOMAS = {
    "Español": "Spanish",
    "Inglés": "English",
    "Francés": "French",
    "Alemán": "German",
    "Portugués": "Portuguese",
    "Italiano": "Italian",
    "Neerlandés": "Dutch",
    "Ruso": "Russian",
    "Chino (simplificado)": "Chinese (simplified)",
    "Japonés": "Japanese",
    "Coreano": "Korean",
    "Árabe": "Arabic",
    "Auto (detección automática)": "auto"
}

# ------------------------------------------------------------
# Funciones de extracción de texto
# ------------------------------------------------------------
def extraer_texto_pdf(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    return "\n".join([pagina.get_text() for pagina in doc])

def extraer_texto_pdf_escaneado(pdf_file):
    imagenes = convert_from_bytes(pdf_file.read())
    texto = ""
    for img in imagenes:
        texto += pytesseract.image_to_string(img, lang='spa+eng') + "\n"
    return texto

def extraer_texto_imagen(imagen_file):
    img = Image.open(imagen_file)
    return pytesseract.image_to_string(img, lang='spa+eng')

def extraer_texto_docx(docx_file):
    doc = Document(docx_file)
    return "\n".join([para.text for para in doc.paragraphs])

def pdf_tiene_texto(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    tiene = any(page.get_text().strip() for page in doc)
    pdf_file.seek(0)
    return tiene

# ------------------------------------------------------------
# División en chunks grandes (10 000 caracteres)
# ------------------------------------------------------------
def dividir_texto(texto, max_chars=10000):
    parrafos = texto.split('\n')
    chunks = []
    actual = ""
    for p in parrafos:
        if len(actual) + len(p) + 1 <= max_chars:
            actual += p + "\n"
        else:
            if actual.strip():
                chunks.append(actual.strip())
            if len(p) > max_chars:
                for i in range(0, len(p), max_chars):
                    chunks.append(p[i:i+max_chars])
                actual = ""
            else:
                actual = p + "\n"
    if actual.strip():
        chunks.append(actual.strip())
    return chunks

# ------------------------------------------------------------
# Traducción con Gemini (contexto solo primer chunk)
# ------------------------------------------------------------
def traducir_chunk(texto, idioma_origen, idioma_destino, contexto=""):
    origen_str = "el idioma detectado automáticamente" if idioma_origen == "auto" else idioma_origen

    if contexto:
        prompt = f"""Eres un traductor profesional. Traduce el siguiente fragmento del {origen_str} al {idioma_destino} de manera **natural y fluida**. Conserva **exactamente** el formato Markdown original (títulos, listas, negritas, cursivas, enlaces). No omitas ningún contenido.
Contexto general del documento: {contexto}

Texto original:
{texto}"""
    else:
        prompt = f"""Eres un traductor profesional. Traduce el siguiente fragmento del {origen_str} al {idioma_destino} de manera **natural y fluida**. Conserva **exactamente** el formato Markdown original (títulos, listas, negritas, cursivas, enlaces). No omitas ningún contenido.

Texto original:
{texto}"""

    respuesta = modelo.generate_content(prompt)
    return respuesta.text

# ------------------------------------------------------------
# Reintentos manuales (solo para rate limit, no para cuota diaria)
# ------------------------------------------------------------
def traducir_con_reintentos(chunk, idioma_origen, idioma_destino, contexto=""):
    max_intentos = 2
    for intento in range(max_intentos):
        try:
            return traducir_chunk(chunk, idioma_origen, idioma_destino, contexto)
        except Exception as e:
            error_msg = str(e)
            # Si es error de cuota diaria, no reintentamos (no servirá)
            if "exceeded your current quota" in error_msg or "429" in error_msg:
                st.error("❌ Has alcanzado el límite diario de peticiones de Gemini. Vuelve a intentarlo mañana o reduce el tamaño del documento.")
                st.stop()
            elif "ResourceExhausted" in error_msg:
                # Rate limit por RPM, reintentamos
                if intento < max_intentos - 1:
                    st.warning(f"⏳ Límite de peticiones por minuto. Esperando 20 segundos...")
                    time.sleep(20)
                else:
                    raise
            else:
                raise

# ------------------------------------------------------------
# Creación de Word (con negritas/cursivas)
# ------------------------------------------------------------
def aplicar_formato(paragraph, texto):
    partes = re.split(r'(\*\*.*?\*\*|\*.*?\*)', texto)
    for parte in partes:
        if parte.startswith('**') and parte.endswith('**'):
            run = paragraph.add_run(parte[2:-2])
            run.bold = True
        elif parte.startswith('*') and parte.endswith('*'):
            run = paragraph.add_run(parte[1:-1])
            run.italic = True
        else:
            paragraph.add_run(parte)

def crear_docx_desde_markdown(markdown):
    doc = Document()
    for linea in markdown.split('\n'):
        linea = linea.strip()
        if not linea:
            doc.add_paragraph('')
            continue
        if linea.startswith('#'):
            nivel = len(re.match(r'^#+', linea).group())
            doc.add_heading(linea.lstrip('#').strip(), level=min(nivel, 9))
        elif linea.startswith('- ') or linea.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            aplicar_formato(p, linea[2:])
        elif re.match(r'^\d+\.\s', linea):
            p = doc.add_paragraph(style='List Number')
            aplicar_formato(p, re.sub(r'^\d+\.\s', '', linea))
        else:
            p = doc.add_paragraph()
            aplicar_formato(p, linea)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ------------------------------------------------------------
# PDF con Helvetica (sin fuentes externas)
# ------------------------------------------------------------
def crear_pdf_desde_markdown(markdown):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    for linea in markdown.split('\n'):
        linea_limpia = linea.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 10, txt=linea_limpia)
    return pdf.output(dest='S')

# ------------------------------------------------------------
# Interfaz de usuario
# ------------------------------------------------------------
st.sidebar.header("⚙️ Configuración de idiomas")
col1, col2 = st.sidebar.columns(2)
with col1:
    idioma_origen_key = st.selectbox(
        "Idioma de origen",
        list(IDIOMAS.keys()),
        index=list(IDIOMAS.keys()).index("Auto (detección automática)")
    )
with col2:
    idioma_destino_key = st.selectbox(
        "Idioma de destino",
        [k for k in IDIOMAS.keys() if k != "Auto (detección automática)"],
        index=0
    )

idioma_origen = IDIOMAS[idioma_origen_key]
idioma_destino = IDIOMAS[idioma_destino_key]

tipo_archivo = st.radio(
    "Tipo de documento a subir:",
    ("PDF (texto digital)", "PDF escaneado o imagen", "Word (.docx)"),
    horizontal=True
)

if tipo_archivo == "Word (.docx)":
    archivo = st.file_uploader("Sube tu archivo", type=["docx"])
else:
    archivo = st.file_uploader("Sube tu archivo", type=["pdf", "png", "jpg", "jpeg"])

# ------------------------------------------------------------
# Procesamiento
# ------------------------------------------------------------
if archivo and st.button("Traducir documento", type="primary"):
    with st.spinner("⏳ Extrayendo texto..."):
        if tipo_archivo == "PDF (texto digital)":
            if archivo.type == "application/pdf":
                if not pdf_tiene_texto(archivo):
                    st.warning("⚠️ Este PDF no contiene texto digital. Cambia a 'PDF escaneado o imagen'.")
                    st.stop()
                archivo.seek(0)
                texto_bruto = extraer_texto_pdf(archivo)
            else:
                st.error("Extensión no soportada.")
                st.stop()
        elif tipo_archivo == "PDF escaneado o imagen":
            if archivo.type == "application/pdf":
                texto_bruto = extraer_texto_pdf_escaneado(archivo)
            else:
                texto_bruto = extraer_texto_imagen(archivo)
        else:
            texto_bruto = extraer_texto_docx(archivo)

        if not texto_bruto.strip():
            st.error("❌ No se pudo extraer texto. El documento podría estar vacío o la imagen es ilegible.")
            st.stop()
        st.success(f"✅ Texto extraído ({len(texto_bruto)} caracteres).")

    # Dividir en chunks grandes
    chunks = dividir_texto(texto_bruto, max_chars=10000)
    st.info(f"📦 Documento dividido en {len(chunks)} fragmento(s).")

    # Mostrar advertencia si el número de peticiones es alto
    if len(chunks) > 100:
        st.warning("⚠️ El documento es muy extenso y puede consumir gran parte de tu cuota diaria. Considera dividirlo en partes.")

    contexto_global = texto_bruto[:300] if len(texto_bruto) > 300 else texto_bruto
    traducciones = []
    progreso = st.progress(0)

    for i, chunk in enumerate(chunks):
        with st.spinner(f"🌍 Traduciendo fragmento {i+1}/{len(chunks)}..."):
            try:
                if i == 0:
                    trad = traducir_con_reintentos(chunk, idioma_origen, idioma_destino, contexto_global)
                else:
                    trad = traducir_con_reintentos(chunk, idioma_origen, idioma_destino)
                traducciones.append(trad)
            except Exception as e:
                st.error(f"❌ Error inesperado al traducir: {str(e)[:200]}")
                st.stop()
        progreso.progress((i + 1) / len(chunks))
        # Pausa de 8 segundos para garantizar < 8 RPM (muy por debajo de las 15)
        time.sleep(8)

    markdown_final = "\n\n".join(traducciones)
    st.success("🎉 ¡Traducción completada!")

    col1, col2 = st.columns(2)
    with col1:
        docx_bytes = crear_docx_desde_markdown(markdown_final)
        st.download_button(
            label="⬇️ DESCARGAR WORD",
            data=docx_bytes,
            file_name="traduccion.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    with col2:
        pdf_bytes = crear_pdf_desde_markdown(markdown_final)
        st.download_button(
            label="⬇️ DESCARGAR PDF",
            data=pdf_bytes,
            file_name="traduccion.pdf",
            mime="application/pdf"
        )
